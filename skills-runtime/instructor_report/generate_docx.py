"""
Instructor Weekly Report Generator
Aggregates student data and creates formatted .docx reports.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import yaml
from pathlib import Path

import structlog

log = structlog.get_logger()


def generate_weekly_report(
    course_id: str,
    instructor_email: str,
    memory_dir: str = "./data/memory",
    output_dir: str = "./data/reports",
) -> str:
    """Generate weekly instructor report as .docx, return file path."""
    stats = aggregate_weekly_stats(course_id, memory_dir)

    doc = Document()

    # Style setup
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("EduClaw Weekly Instructor Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Course: {stats['course_name']}")
    doc.add_paragraph(f"Period: {stats['period_start']} to {stats['period_end']}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Instructor: {instructor_email}")
    doc.add_paragraph(f"Total active students: {stats['active_students']}")

    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(
        f"This week, {stats['total_doubts']} student doubts were recorded. "
        f"The most confused topic was '{stats['top_weak_topic']}' with "
        f"{stats['top_weak_count']} questions. Average quiz score: "
        f"{stats['avg_quiz_score']:.1f}%."
    )

    # Top weak topics table
    doc.add_heading("Top 5 Topics With Most Doubts", 1)
    if stats["weak_topics"]:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Topic"
        hdr[1].text = "Doubt Count"
        hdr[2].text = "Avg Quiz Score"

        for topic, count, score in stats["weak_topics"][:5]:
            row = table.add_row().cells
            row[0].text = topic
            row[1].text = str(count)
            row[2].text = f"{score:.1f}%"
    else:
        doc.add_paragraph("No doubts recorded this week.")

    # Recommended re-teaching
    doc.add_heading("Recommended Re-Teaching Topics", 1)
    if stats["reteach_recommendations"]:
        for topic in stats["reteach_recommendations"]:
            doc.add_paragraph(f"• {topic}", style="List Bullet")
    else:
        doc.add_paragraph("All topics performing well — no re-teaching needed.")

    # Save
    output_path = (
        Path(output_dir)
        / f"{course_id}_report_{datetime.now().strftime('%Y%m%d')}.docx"
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    log.info("report_generated", course_id=course_id, path=str(output_path))
    return str(output_path)


def aggregate_weekly_stats(course_id: str, memory_dir: str) -> dict:
    """Aggregate all student YAML files for the past 7 days."""
    students_dir = Path(memory_dir) / "students"
    week_ago = datetime.now() - timedelta(days=7)

    doubt_counts: dict[str, int] = {}
    quiz_scores: dict[str, list[int]] = {}
    active_students = 0
    total_doubts = 0

    if students_dir.exists():
        for yaml_file in students_dir.glob("*.yaml"):
            with open(yaml_file) as f:
                profile = yaml.safe_load(f)

            if not profile or course_id not in profile.get("enrolled_courses", []):
                continue

            active_students += 1

            # Count doubts this week
            for doubt in profile.get("doubt_log", []):
                try:
                    doubt_date = datetime.fromisoformat(doubt["date"])
                except (ValueError, KeyError):
                    continue
                if doubt_date >= week_ago and doubt.get("course") == course_id:
                    total_doubts += 1
                    topic_key = " ".join(doubt["question"].lower().split()[:3])
                    doubt_counts[topic_key] = doubt_counts.get(topic_key, 0) + 1

            # Collect quiz scores
            for topic, scores in (
                profile.get("quiz_scores", {}).get(course_id, {}).items()
            ):
                quiz_scores.setdefault(topic, []).extend(scores)

    # Sort by doubt count
    weak_topics_sorted = sorted(doubt_counts.items(), key=lambda x: x[1], reverse=True)

    # Compute avg quiz scores per topic
    topic_with_scores = [
        (
            topic,
            count,
            sum(quiz_scores.get(topic.replace(" ", "_"), [50]))
            / max(1, len(quiz_scores.get(topic.replace(" ", "_"), [50]))),
        )
        for topic, count in weak_topics_sorted
    ]

    # Reteach: topics with avg score < 60%
    reteach = [t for t, c, s in topic_with_scores if s < 60.0]

    all_scores = [s for scores in quiz_scores.values() for s in scores]
    avg_score = sum(all_scores) / max(1, len(all_scores)) if all_scores else 0.0

    return {
        "course_name": course_id.replace("_", " ").title(),
        "period_start": week_ago.strftime("%Y-%m-%d"),
        "period_end": datetime.now().strftime("%Y-%m-%d"),
        "active_students": active_students,
        "total_doubts": total_doubts,
        "weak_topics": topic_with_scores,
        "top_weak_topic": weak_topics_sorted[0][0] if weak_topics_sorted else "N/A",
        "top_weak_count": weak_topics_sorted[0][1] if weak_topics_sorted else 0,
        "avg_quiz_score": avg_score,
        "reteach_recommendations": reteach[:3],
    }
